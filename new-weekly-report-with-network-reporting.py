import os
import requests
import json
import time
from requests.auth import HTTPBasicAuth
from docx import Document
from docx.shared import Inches
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
import smtplib
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
from email.mime.multipart import MIMEMultipart
from datetime import datetime

# Get Timestamp
timestamp = datetime.now().strftime('%Y-%m-%d')
# [Your Zabbix API and Grafana credentials and functions are here...]
# Zabbix API endpoint
ZABBIX_SERVER = "http://{abstracted}/zabbix/api_jsonrpc.php"
USERNAME = zabbix_username
PASSWORD = zabbix_password

# Grafana credentials
GRAFANA_USERNAME = grafana_username
GRAFANA_PASSWORD = grafana_password

# Time Range (Past Week)
current_time = int(time.time())
one_week_ago = current_time - (7 * 24 * 60 * 60)

# Authenticate with Zabbix API
headers = {
    "Content-Type": "application/json-rpc"
}
payload = {
    "jsonrpc": "2.0",
    "method": "user.login",
    "params": {
        "user": USERNAME,
        "password": PASSWORD
    },
    "id": 1,
    "auth": None
}
response = requests.post(ZABBIX_SERVER, headers=headers, data=json.dumps(payload))
auth_token = response.json()["result"]

# Function to get history data
def get_history_data(item_id, time_from, time_till):
    payload = {
        "jsonrpc": "2.0",
        "method": "history.get",
        "params": {
            "output": ["itemid", "clock", "value"],
            "itemids": [item_id],
            "time_from": time_from,
            "time_till": time_till,
            "sortfield": "clock",
            "sortorder": "DESC"
        },
        "auth": auth_token,
        "id": 3
    }
    response = requests.post(ZABBIX_SERVER, headers=headers, data=json.dumps(payload))
    return response.json()["result"]

# Function to send an email
def send_email(subject, body, to, file_path):
    # Email credentials
    email_sender = email_sender
    email_password = email_password
    
    msg = MIMEMultipart()
    msg['From'] = email_sender
    msg['To'] = ', '.join(to)
    msg['Subject'] = subject
    
    msg.attach(MIMEText(body, "plain"))
    
    with open(file_path, "rb") as f:
        attach = MIMEApplication(f.read(),_subtype="docx")
        attach.add_header('Content-Disposition','attachment',filename=str(file_path))
        msg.attach(attach)

    server = smtplib.SMTP('smtp.gmail.com', 587)
    server.starttls()
    server.login(email_sender, email_password)
    server.send_message(msg)
    server.close()
################GET NETWORK AVAILABILITY##########################################################################
# Fetch network availability data for the "Network Monitoring" group
network_availability_data = {}
group_name = "Network Monitoring"
payload = {
    "jsonrpc": "2.0",
    "method": "host.get",
    "params": {
        "output": ["hostid", "name"],
        "groupids": "30"
    },
    "auth": auth_token,
    "id": 2
}
response = requests.post(ZABBIX_SERVER, headers=headers, data=json.dumps(payload))
hosts = response.json()["result"]

for host in hosts:
    host_id = host["hostid"]
    host_name = host["name"]
    # Strip the " Network Monitoring" suffix to get the corresponding device group name
    device_group_name = host_name.replace(" Network Monitoring", "")
    
    # Get the item ID for ICMP ping for the host
    payload = {
        "jsonrpc": "2.0",
        "method": "item.get",
        "params": {
            "output": ["itemid"],
            "hostids": host_id,
            "search": {
                "key_": "icmpping"
            }
        },
        "auth": auth_token,
        "id": 4
    }
    response = requests.post(ZABBIX_SERVER, headers=headers, data=json.dumps(payload))
    ITEM_ID_ICMP_PING = response.json()["result"][0]["itemid"]

    # Fetch ICMP ping data
    icmp_ping_data = get_history_data(ITEM_ID_ICMP_PING, one_week_ago, current_time)
    total_pings = len(icmp_ping_data)
    successful_pings = sum(int(data_point["value"]) == 1 for data_point in icmp_ping_data)
    availability_percentage = (successful_pings / total_pings) * 100 if total_pings > 0 else 0
    network_availability_data[device_group_name] = availability_percentage
#######################################################################################################################


# [Your Zabbix API and Grafana code here...]
# Define device groups
device_groups = ["C2 VAL HE03-SG03101 L-CHTR QA02 Canary Monitoring", "C2 VAL HE11-SG11201 L-CHTR QA02 Canary Monitoring", "C2 VAL HE20-SG20101 L-CHTR UAT Canary Monitoring", "C2 VAL HE23-H1-SG1 L-CHTR UAT Canary Monitoring", "C2 VAL HE24-SG24101 L-CHTR UAT Canary Monitoring", "SP HE03-SG03101 L-CHTR QA02 Canary Monitoring", "SP HE20-SG20102 L-CHTR UAT Canary Monitoring", "C2 VAL HE02-SG02118 L-TWC UAT Canary Monitoring", "C2 VAL HE05-SG05101 L-TWC QA02 Canary Monitoring"]

# URLs for Modem Grafana image retrieval
urls = {
    "C2 VAL HE03-SG03101 L-CHTR QA02 Canary Monitoring": "http://{abstracted}:3000/render/d-solo/ddfe859f-5b71-4710-8c2b-b9c608b70cf2/80f3d428-4985-5e2d-8d91-433e599f2e11",
    "C2 VAL HE02-SG02118 L-TWC UAT Canary Monitoring": "http://{abstracted}:3000/render/d-solo/ef70b10c-bcd3-47c7-9fec-d56b49729bca/d6b2dba1-a82f-5ab8-819d-d84c23d9d63d",
    "C2 VAL HE05-SG05101 L-TWC QA02 Canary Monitoring": "http://{abstracted}:3000/render/d-solo/eb766426-ae20-4d0b-8043-ab05f950a032/c4a4da37-918e-5ec8-8b88-dacae8a62c60",
    "C2 VAL HE11-SG11201 L-CHTR QA02 Canary Monitoring": "http://{abstracted}:3000/render/d-solo/ddfe859f-5b71-4710-8c2b-b9c608b70cf3/d021b2f6-0d20-57c0-9269-628ba35f3a6e",
    "C2 VAL HE20-SG20101 L-CHTR UAT Canary Monitoring": "http://{abstracted}:3000/render/d-solo/a2932155-5555-4b0a-af8e-262e6edc784a/c615a6d9-93a6-5068-9245-469c74fcf0e2",
    "C2 VAL HE23-H1-SG1 L-CHTR UAT Canary Monitoring": "http://{abstracted}:3000/render/d-solo/a98f25c8-4462-48b1-a612-df359f2e1326/f75e501d-6916-5115-83ce-fbe4fbbb43f8",
    "C2 VAL HE24-SG24101 L-CHTR UAT Canary Monitoring": "http://{abstracted}:3000/render/d-solo/cc8e2411-0f96-475a-972c-6a5f50757c66/0caebbad-ff6d-5e15-8895-62b2cd30e0e3",
    "SP HE03-SG03101 L-CHTR QA02 Canary Monitoring": "http://{abstracted}:3000/render/d-solo/b993c393-e2b5-44f6-b234-0b1a98f349ff/67dfe480-1477-52f0-8c1c-f3d1fbbb28a0",
    "SP HE20-SG20102 L-CHTR UAT Canary Monitoring": "http://{abstracted}:3000/render/d-solo/b6c6aff3-2a2d-4106-9536-d89e363528da/da9318b0-0248-5dbe-b641-d091efb285d0"
  }

# URLs for Network health Grafana image retrieval
network_urls = {'C2 VAL HE03-SG03101 L-CHTR QA02 Canary Monitoring': 'http://{abstracted}:3000/render/d-solo/e8e130ce-1934-4923-92b1-5d38a32d6568?orgId=1&panelId=1&width=1000&height=500&tz=America%2FDenver', 'C2 VAL HE02-SG02118 L-TWC UAT Canary Monitoring': 'http://192.219.218.106:3000/render/d-solo/d26a681a-57e6-4942-ba52-7dbf24ba71d3?orgId=1&panelId=2&width=1000&height=500&tz=America%2FDenver', 'C2 VAL HE05-SG05101 L-TWC QA02 Canary Monitoring': 'http://{abstracted}:3000/render/d-solo/cb0af2ad-cf1a-4a39-8d41-8ba86c76e601?orgId=1&panelId=1&width=1000&height=500&tz=America%2FDenver', 'C2 VAL HE11-SG11201 L-CHTR QA02 Canary Monitoring': 'http://{abstracted}:3000/render/d-solo/e04b025c-1fee-4150-8512-3d4c5411500b?orgId=1&panelId=1&width=1000&height=500&tz=America%2FDenver', 'C2 VAL HE20-SG20101 L-CHTR UAT Canary Monitoring': 'http://{abstracted}:3000/render/d-solo/a1b60e90-036d-4be1-94a2-39651161749e?orgId=1&panelId=1&width=1000&height=500&tz=America%2FDenver', 'C2 VAL HE23-H1-SG1 L-CHTR UAT Canary Monitoring': 'http://{abstracted}:3000/render/d-solo/f4ca7883-e1ba-47ad-9807-fbb47cd64778?orgId=1&panelId=1&width=1000&height=500&tz=America%2FDenver', 'C2 VAL HE24-SG24101 L-CHTR UAT Canary Monitoring': 'http://{abstracted}:3000/render/d-solo/c19ea2fb-b570-44ba-9d6a-dbb1c2357496?orgId=1&panelId=1&width=1000&height=500&tz=America%2FDenver', 'SP HE03-SG03101 L-CHTR QA02 Canary Monitoring': 'http://{abstracted}:3000/render/d-solo/c944014d-a812-4cda-a392-e41dbe8c4b75?orgId=1&panelId=1&width=1000&height=500&tz=America%2FDenver', 'SP HE20-SG20102 L-CHTR UAT Canary Monitoring': 'http://192.219.218.106:3000/render/d-solo/ad8dacba-a41e-4412-8769-351fe9bdc476?orgId=1&panelId=1&width=1000&height=500&tz=America%2FDenver'}

document = Document()
# Creating a Word document
# Define styles for headings and paragraphs
styles = document.styles

heading1_style = styles.add_style('CustomHeading1', WD_STYLE_TYPE.PARAGRAPH)
heading1_style.base_style = styles['Heading 1']
heading1_style.font.name = 'Arial'
heading1_style.font.size = Pt(16)
heading1_style.font.bold = True

heading2_style = styles.add_style('CustomHeading2', WD_STYLE_TYPE.PARAGRAPH)
heading2_style.base_style = styles['Heading 2']
heading2_style.font.name = 'Arial'
heading2_style.font.size = Pt(14)
heading2_style.font.bold = True

normal_style = styles.add_style('CustomNormal', WD_STYLE_TYPE.PARAGRAPH)
normal_style.base_style = styles['Normal']
normal_style.font.name = 'Arial'
normal_style.font.size = Pt(12)

document.add_heading(f'Weekly Modem Availability Report {timestamp}', 0).style = heading1_style

for group_name in device_groups:
    total_availability = 0
    total_devices = 0
    document.add_heading(group_name, level=1).style = heading2_style  # Add the device group label to the report
    
    # [Code to get group_id and hosts as per your provided script...]
    print(f"\nDevice Group: {group_name}")

    # Get the group ID for the group name
    payload = {
        "jsonrpc": "2.0",
        "method": "hostgroup.get",
        "params": {
            "output": ["groupid"],
            "filter": {
                "name": [group_name]
            }
        },
        "auth": auth_token,
        "id": 2
    }
    response = requests.post(ZABBIX_SERVER, headers=headers, data=json.dumps(payload))
    group_id = response.json()["result"][0]["groupid"]

    # Get host IDs in the device group
    payload = {
        "jsonrpc": "2.0",
        "method": "host.get",
        "params": {
            "output": ["hostid", "name"],
            "groupids": group_id
        },
        "auth": auth_token,
        "id": 3
    }
    response = requests.post(ZABBIX_SERVER, headers=headers, data=json.dumps(payload))
    hosts = response.json()["result"]
    panel_id = 2
    # Prepare table data
    table_data = []
    network_availability = network_availability_data.get(group_name, None)
    network_availability = network_availability_data.get(group_name, None)
    if network_availability is not None:
        # Fetch the Grafana graph for network availability
        base_url = network_urls.get(group_name)
        if base_url:
            response = requests.get(base_url, auth=HTTPBasicAuth(GRAFANA_USERNAME, GRAFANA_PASSWORD))
            if response.status_code == 200:
                graph_path = f"{group_name.replace(' ', '_')}_Network_Monitoring.png"
                with open(graph_path, "wb") as f:
                    f.write(response.content)
                
                # Append the group name, network availability, and image path to the table data
                table_data.append((f"{group_name} Network Availability", network_availability, graph_path))
    
    for host in sorted(hosts, key=lambda x: x["name"]):
        host_id = host["hostid"]
        host_name = host["name"]
        
        # [Code to get ITEM_ID_ICMP_PING and icmp_ping_data as per your provided script...]
        # Get the item ID for ICMP ping for the host
        payload = {
            "jsonrpc": "2.0",
            "method": "item.get",
            "params": {
                "output": ["itemid"],
                "hostids": host_id,
                "search": {
                    "key_": "icmpping"
                }
            },
            "auth": auth_token,
            "id": 4
        }
        response = requests.post(ZABBIX_SERVER, headers=headers, data=json.dumps(payload))
        ITEM_ID_ICMP_PING = response.json()["result"][0]["itemid"]

        # Fetch ICMP ping data
        icmp_ping_data = get_history_data(ITEM_ID_ICMP_PING, one_week_ago, current_time)


        total_pings = len(icmp_ping_data)
        successful_pings = sum(int(data_point["value"]) == 1 for data_point in icmp_ping_data)
        availability_percentage = (successful_pings / total_pings) * 100 if total_pings > 0 else 0
        # Store availability_percentages in a list
        total_availability += availability_percentage
        total_devices += 1
        # Store table data
        host_name = host["name"]

        
        base_url = urls.get(group_name)
        if base_url:
            params = {
                "orgId": 1,
                "panelId": panel_id,
                "width": 1000,
                "height": 500,
                "tz": "America/Denver"
            }
            response = requests.get(base_url, params=params, auth=HTTPBasicAuth(GRAFANA_USERNAME, GRAFANA_PASSWORD))
            if response.status_code == 200:
                file_path = os.path.join(group_name.replace(" ", "_"), f"{host_name.replace(' ', '_')}.png")
                os.makedirs(os.path.dirname(file_path), exist_ok=True)
                with open(file_path, "wb") as f:
                    f.write(response.content)
                
                # Add image to the Word document
                table_data.append((host_name, availability_percentage, file_path))

        
        panel_id += 1
        
    # Display average availability
    average_availability = total_availability / total_devices if total_devices > 0 else 0
    document.add_paragraph(f"Average availability for {group_name}: {average_availability:.2f}%").style = normal_style

    # Add a table to the document
    table = document.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    # Add header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Device Name'
    header_cells[1].text = 'Availability'
    header_cells[2].text = 'Graph'
    
    # Add data rows
    for host_name, availability_percentage, graph_file_path in table_data:
        row_cells = table.add_row().cells
        row_cells[0].text = host_name
        row_cells[1].text = f"1 week: {availability_percentage:.2f}%"
        if graph_file_path:
            row_cells[2].paragraphs[0].add_run().add_picture(graph_file_path, width=Inches(2))

# Save the document
report_filename = "Weekly_Modem_Availability_Report.docx"
document.save(report_filename)

# Send the document via email
subject = "Weekly Modem Availability Report"
body = "Please find attached the Weekly Modem Availability Report."
send_email(subject, body, recipient_email, report_filename)

print("Report sent.")
